"""
document_extractor.py — Text extraction + client fact extraction from uploaded documents.

Supported file types:
  application/pdf           → pdfplumber (text layer)
  application/vnd...docx   → python-docx (paragraphs)
  image/png, jpeg, webp    → Azure OpenAI vision (GPT-4o mini)

Client fact extraction:
  Uses Azure OpenAI with the same canonical memory schema as memory_extractor.py.
  Returns a delta dict ready to be passed to merge_delta().
  Returns {} on any failure (never raises).
"""

from __future__ import annotations

import base64
import json
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def _normalize_to_canonical_schema(extracted: dict) -> dict:
    """
    Normalize document-extracted keys to the canonical memory schema used by
    memory_extractor + build_tool_input_from_memory.
    """
    normalized = dict(extracted)

    personal = normalized.get("personal")
    if isinstance(personal, dict):
        if "smoker" in personal and "is_smoker" not in personal:
            personal["is_smoker"] = personal.pop("smoker")
        # Some documents report dependant count under financial; keep canonical
        # count on personal.dependants for downstream tool mapping.
        normalized["personal"] = personal

    financial = normalized.get("financial")
    if isinstance(financial, dict):
        if "annual_income" in financial and "annual_gross_income" not in financial:
            financial["annual_gross_income"] = financial.pop("annual_income")
        if "number_of_dependants" in financial:
            dependant_count = financial.pop("number_of_dependants")
            p = normalized.setdefault("personal", {})
            if isinstance(p, dict) and "dependants" not in p:
                p["dependants"] = dependant_count
        normalized["financial"] = financial

    insurance = normalized.get("insurance")
    if isinstance(insurance, dict):
        if "insurer" in insurance and "insurer_name" not in insurance:
            insurance["insurer_name"] = insurance.pop("insurer")

        # Map common document cover amounts into canonical insured amounts.
        if "existing_life_cover" in insurance and "life_sum_insured" not in insurance:
            insurance["life_sum_insured"] = insurance.pop("existing_life_cover")
        if "existing_tpd_cover" in insurance and "tpd_sum_insured" not in insurance:
            insurance["tpd_sum_insured"] = insurance.pop("existing_tpd_cover")
        if "existing_income_protection_cover" in insurance and "ip_monthly_benefit" not in insurance:
            insurance["ip_monthly_benefit"] = insurance.pop("existing_income_protection_cover")
        if "existing_trauma_cover" in insurance and "trauma_sum_insured" not in insurance:
            insurance["trauma_sum_insured"] = insurance.pop("existing_trauma_cover")

        # If any existing cover amount is present, infer existing policy flag.
        has_any_cover = any(
            insurance.get(k) is not None
            for k in ("life_sum_insured", "tpd_sum_insured", "ip_monthly_benefit", "trauma_sum_insured")
        )
        if has_any_cover and "has_existing_policy" not in insurance:
            insurance["has_existing_policy"] = True
        normalized["insurance"] = insurance

    goals = normalized.get("goals")
    if isinstance(goals, dict) and "wants_inside_super" in goals:
        wants_inside_super = goals.pop("wants_inside_super")
        ins = normalized.setdefault("insurance", {})
        if isinstance(ins, dict) and "in_super" not in ins:
            ins["in_super"] = wants_inside_super
        normalized["goals"] = goals

    return normalized

# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_text_from_pdf(file_path: str) -> str:
    """Extract all text from a PDF using pdfplumber."""
    try:
        import pdfplumber
        pages: list[str] = []
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"[Page {i + 1}]\n{text}")
        return "\n\n".join(pages)
    except Exception as exc:
        logger.warning("PDF extraction failed for %s: %s", file_path, exc)
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """Extract all paragraphs from a DOCX file."""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n".join(paragraphs)
    except Exception as exc:
        logger.warning("DOCX extraction failed for %s: %s", file_path, exc)
        return ""


async def extract_text_from_image(file_path: str) -> str:
    """
    Use Azure OpenAI vision to describe an image and extract any visible text.
    Falls back to empty string if the model doesn't support vision.
    """
    try:
        from app.core.llm import get_chat_model_fresh

        with open(file_path, "rb") as f:
            image_data = base64.b64encode(f.read()).decode("utf-8")

        suffix = Path(file_path).suffix.lower()
        mime_map = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".webp": "image/webp"}
        mime = mime_map.get(suffix, "image/png")

        llm = get_chat_model_fresh(temperature=0.0)

        messages = [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "This image is from an insurance advisory application. "
                            "Please extract ALL visible text from this image verbatim, "
                            "preserving structure (tables, lists, headings). "
                            "Then briefly describe any charts or non-text elements."
                        ),
                    },
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:{mime};base64,{image_data}"},
                    },
                ],
            }
        ]

        response = await llm.ainvoke(messages)
        return response.content.strip() if hasattr(response, "content") else ""
    except Exception as exc:
        logger.warning("Image extraction failed for %s: %s", file_path, exc)
        return ""


async def extract_text(file_path: str, content_type: str) -> str:
    """
    Route to the correct extractor based on content_type.
    Returns extracted text (may be empty string on failure).
    """
    if content_type == "application/pdf":
        return extract_text_from_pdf(file_path)
    if content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_text_from_docx(file_path)
    if content_type.startswith("image/"):
        return await extract_text_from_image(file_path)
    logger.warning("Unsupported content_type for text extraction: %s", content_type)
    return ""


# ---------------------------------------------------------------------------
# Client fact extraction (canonical memory schema)
# ---------------------------------------------------------------------------

_FACT_EXTRACTION_SYSTEM_PROMPT = """\
You are a data-extraction assistant for an insurance advisory platform.
Read the document text and extract client facts in the exact JSON structure below.
Only extract facts explicitly stated — do NOT infer or guess.
Return ONLY a JSON object. No markdown, no explanation.

Schema (all fields optional — omit if not found):
{
  "personal": {
    "full_name": null,
    "age": null,
    "date_of_birth": null,
    "gender": null,
    "occupation": null,
    "employment_status": null,
    "smoker": null,
    "state_of_residence": null
  },
  "financial": {
    "annual_income": null,
    "monthly_income": null,
    "total_assets": null,
    "total_liabilities": null,
    "net_worth": null,
    "super_balance": null,
    "monthly_expenses": null,
    "number_of_dependants": null
  },
  "insurance": {
    "existing_life_cover": null,
    "existing_tpd_cover": null,
    "existing_income_protection_cover": null,
    "existing_trauma_cover": null,
    "cover_types": [],
    "policy_owner": null,
    "policy_number": null,
    "insurer": null,
    "premium_frequency": null,
    "annual_premium": null
  },
  "health": {
    "height_cm": null,
    "weight_kg": null,
    "medical_conditions": [],
    "current_medications": [],
    "hazardous_activities": [],
    "family_history": null
  },
  "goals": {
    "primary_goal": null,
    "risk_tolerance": null,
    "retirement_age": null,
    "years_to_retirement": null,
    "wants_inside_super": null
  },
  "_meta": {
    "document_type": null,
    "document_date": null,
    "confidence_note": null
  }
}

Rules:
- Set numeric fields as numbers (not strings): age=42, annual_income=120000
- Set boolean fields as true/false
- List fields: use [] if none found, or a list of strings
- _meta.document_type: classify as one of: "policy_schedule", "soa", "client_profile", "medical_report", "payslip", "tax_return", "other"
- _meta.confidence_note: brief note on extraction quality (e.g. "clear SOA with full client details")
- Omit any top-level section entirely if no fields within it were found
"""


async def extract_client_facts(text: str, filename: str) -> dict:
    """
    Use Azure OpenAI to extract canonical client facts from extracted document text.
    Returns a delta dict (same format as memory_extractor output).
    Returns {} on any failure.
    """
    if not text or not text.strip():
        return {}

    try:
        from app.core.llm import get_chat_model_fresh

        llm = get_chat_model_fresh(temperature=0.0)

        # Limit text to 6000 chars to stay within context; take from start + end
        if len(text) > 6000:
            text_sample = text[:4000] + "\n...[middle truncated]...\n" + text[-2000:]
        else:
            text_sample = text

        user_prompt = f"Document filename: {filename}\n\n{text_sample}"

        messages = [
            {"role": "system", "content": _FACT_EXTRACTION_SYSTEM_PROMPT},
            {"role": "user",   "content": user_prompt},
        ]

        response = await llm.ainvoke(messages)
        raw = response.content.strip() if hasattr(response, "content") else ""

        # Strip markdown fences if present
        if raw.startswith("```"):
            lines = [l for l in raw.splitlines() if not l.strip().startswith("```")]
            raw = "\n".join(lines).strip()

        data = json.loads(raw)
        if not isinstance(data, dict):
            return {}

        # Remove top-level sections that are entirely null/empty
        cleaned: dict = {}
        for section, values in data.items():
            if section == "_meta":
                if any(v for v in values.values() if v is not None):
                    cleaned["_meta"] = values
                continue
            if not isinstance(values, dict):
                continue
            non_null = {k: v for k, v in values.items() if v is not None and v != [] and v != ""}
            if non_null:
                cleaned[section] = non_null

        cleaned = _normalize_to_canonical_schema(cleaned)

        logger.info(
            "document_extractor: extracted facts from '%s' — sections: %s",
            filename,
            list(cleaned.keys()),
        )
        return cleaned

    except Exception as exc:
        logger.warning("extract_client_facts failed for '%s': %s", filename, exc)
        return {}


def facts_summary(extracted_facts: dict) -> str:
    """
    Build a short human-readable summary of extracted facts for the upload response.
    """
    if not extracted_facts:
        return "No client facts detected."

    parts: list[str] = []

    personal = extracted_facts.get("personal", {})
    if personal.get("full_name"):
        parts.append(f"Client: {personal['full_name']}")
    if personal.get("age"):
        parts.append(f"Age: {personal['age']}")
    if personal.get("occupation"):
        parts.append(f"Occupation: {personal['occupation']}")

    financial = extracted_facts.get("financial", {})
    annual_income = financial.get("annual_gross_income", financial.get("annual_income"))
    if annual_income:
        parts.append(f"Annual income: ${annual_income:,.0f}")
    if financial.get("super_balance"):
        parts.append(f"Super balance: ${financial['super_balance']:,.0f}")

    insurance = extracted_facts.get("insurance", {})
    insurer = insurance.get("insurer_name", insurance.get("insurer"))
    if insurer:
        parts.append(f"Insurer: {insurer}")
    if insurance.get("policy_number"):
        parts.append(f"Policy #: {insurance['policy_number']}")

    meta = extracted_facts.get("_meta", {})
    doc_type = meta.get("document_type")

    header = f"[{doc_type.upper()}] " if doc_type else ""
    return header + ("; ".join(parts) if parts else "Document processed — no specific client facts extracted.")
